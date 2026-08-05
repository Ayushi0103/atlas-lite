from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path
from shutil import copyfileobj
from uuid import uuid4
import logging

from docx import Document as DocxDocument
from fastapi import FastAPI, File, HTTPException, UploadFile, status
from pydantic import BaseModel
from pypdf import PdfReader
from sqlmodel import Session, select, or_

from app.connectors.youtube import (
    InvalidYouTubeUrlError,
    YouTubeTranscriptUnavailableError,
    YouTubeVideoUnavailableError,
    download_transcript,
    extract_video_id,
    normalize_source_url,
)
from app.database import SessionDep, create_db_and_tables
from app.models import Collection, CollectionNote, Document, Note
from app.routes.ai import router as ai_router
from app.routes.conversations import router as conversations_router
from app.routes.search import router as search_router
from app.services.embedding import (
    add_document_embedding,
    add_note_embedding,
    delete_embedding,
    update_embedding,
)
from app.services.ocr import OCRExtractionError, extract_text_from_image
from app.services.summary_service import (
    build_document_embedding_text,
    summarize_and_store_document_metadata,
)

logger = logging.getLogger(__name__)

ROOT_DIR = Path(__file__).resolve().parents[2]
UPLOAD_DIR = ROOT_DIR / "uploads"
SUPPORTED_IMAGE_TYPES = {"png", "jpg", "jpeg", "webp"}
SUPPORTED_DOCUMENT_TYPES = {"txt", "pdf", "docx", "md", *SUPPORTED_IMAGE_TYPES}


@asynccontextmanager
async def lifespan(app: FastAPI):
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    create_db_and_tables()
    yield


app = FastAPI(lifespan=lifespan)


class NoteCreate(BaseModel):
    title: str
    content: str
    tags: str


class CollectionCreate(BaseModel):
    name: str
    description: str | None = None


class YouTubeImportRequest(BaseModel):
    url: str


def safely_index_note(note: Note) -> None:
    if note.id is None:
        return

    try:
        add_note_embedding(note.id, note.title, note.content)
    except Exception:
        logger.exception("Failed to index note %s", note.id)


def safely_index_document(document: Document) -> None:
    if document.id is None:
        return

    try:
        add_document_embedding(
            document.id,
            document.filename,
            build_document_embedding_text(document),
        )
    except Exception:
        logger.exception("Failed to index document %s", document.id)


def safely_summarize_document(document: Document, session: Session) -> None:
    if document.id is None:
        return

    try:
        summarize_and_store_document_metadata(document, session)
        safely_index_document(document)
    except Exception:
        session.rollback()
        logger.exception("Failed to summarize document %s", document.id)


def safely_update_note_embedding(note: Note) -> None:
    if note.id is None:
        return

    try:
        update_embedding(
            source_type="note",
            source_id=note.id,
            title=note.title,
            text=f"{note.title}\n{note.content}",
        )
    except Exception:
        logger.exception("Failed to update embedding for note %s", note.id)


def safely_delete_embedding(source_type: str, source_id: int | None) -> None:
    if source_id is None:
        return

    try:
        delete_embedding(source_type, source_id)  # type: ignore[arg-type]
    except Exception:
        logger.exception("Failed to delete embedding for %s %s", source_type, source_id,)


def get_document_file_type(filename: str) -> str:
    file_type = Path(filename).suffix.lower().lstrip(".")
    if file_type not in SUPPORTED_DOCUMENT_TYPES:
        raise HTTPException(
            status_code=400,
            detail=(
                "Unsupported file type. Upload a TXT, MD, PDF, DOCX, PNG, JPG, "
                "JPEG, or WEBP file."
            ),
        )

    return file_type


def save_upload_file(upload_file: UploadFile) -> Path:
    safe_filename = Path(upload_file.filename or "").name
    if not safe_filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    saved_filename = f"{uuid4().hex}_{safe_filename}"
    destination = UPLOAD_DIR / saved_filename

    with destination.open("wb") as buffer:
        copyfileobj(upload_file.file, buffer)

    return destination


def extract_text_from_document(file_path: Path, file_type: str) -> str:
    if file_type in SUPPORTED_IMAGE_TYPES:
        text = extract_text_from_image(str(file_path))
        if not text.strip():
            raise OCRExtractionError("Could not extract readable text from image.")

        return text

    if file_type in {"txt", "md"}:
        return file_path.read_text(encoding="utf-8", errors="replace")

    if file_type == "pdf":
        reader = PdfReader(str(file_path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if file_type == "docx":
        doc = DocxDocument(str(file_path))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    raise HTTPException(status_code=400, detail="Unsupported file type")


@app.get("/")
def read_root():
    return {"message": "Welcome to Atlas Lite"}


@app.post("/notes")
def save_note(note: NoteCreate, session: SessionDep):
    new_note = Note(title=note.title, content=note.content, tags=note.tags)
    session.add(new_note)
    session.commit()
    session.refresh(new_note)
    safely_index_note(new_note)

    return {
        "status": "saved",
        "note": new_note,
    }


@app.get("/notes")
def get_notes(session: SessionDep):
    return session.exec(select(Note)).all()


# Register a GET endpoint for searching notes.
@app.get("/notes/search")
def search_notes(q: str, session: SessionDep):
    q = q.strip()

    if not q:
        raise HTTPException(
            status_code=400,
            detail="Search query cannot be empty",
        )

    statement = select(Note).where(
        or_(
            Note.title.contains(q),
            Note.content.contains(q),
            Note.tags.contains(q),
        )
    )

    return session.exec(statement).all()


@app.get("/notes/{note_id}")
def get_note(note_id: int, session: SessionDep):
    note = session.get(Note, note_id)
    if note is None:
        raise HTTPException(status_code=404, detail="Note not found")

    return note


@app.post("/documents", status_code=status.HTTP_201_CREATED)
def upload_document(session: SessionDep, file: UploadFile = File(...)):
    if file.filename is None:
        raise HTTPException(status_code=400, detail="Filename is required")

    file_type = get_document_file_type(file.filename)
    saved_path = save_upload_file(file)

    try:
        text_content = extract_text_from_document(saved_path, file_type)
    except OCRExtractionError as exc:
        saved_path.unlink(missing_ok=True)
        raise HTTPException(
            status_code=400,
            detail="Could not extract readable text from image.",
        ) from exc
    except Exception as exc:
        saved_path.unlink(missing_ok=True)
        raise HTTPException(
            status_code=400, detail=f"Could not extract text from document: {exc}"
        ) from exc

    document = Document(
        filename=Path(file.filename).name,
        file_type=file_type,
        file_path=str(saved_path.relative_to(ROOT_DIR)),
        text_content=text_content,
    )
    session.add(document)
    session.commit()
    session.refresh(document)
    safely_summarize_document(document, session)

    return {
        "status": "saved",
        "document": document,
    }


@app.get("/documents")
def get_documents(session: SessionDep):
    statement = select(Document).order_by(Document.created_at.desc())
    return session.exec(statement).all()


@app.get("/documents/search")
def search_documents(q: str, session: SessionDep):
    q = q.strip()

    if not q:
        raise HTTPException(
            status_code=400,
            detail="Search query cannot be empty",
        )

    statement = select(Document).where(
        or_(
            Document.filename.contains(q),
            Document.text_content.contains(q),
            Document.short_summary.contains(q),
            Document.detailed_summary.contains(q),
            Document.key_concepts.contains(q),
            Document.keywords.contains(q),
            Document.suggested_questions.contains(q),
        )
    )

    return session.exec(statement).all()


@app.get("/documents/{document_id}")
def get_document(document_id: int, session: SessionDep):
    document = session.get(Document, document_id)
    if document is None:
        raise HTTPException(status_code=404, detail="Document not found")

    return document


@app.post("/connectors/youtube", status_code=status.HTTP_201_CREATED)
def import_youtube_transcript(request: YouTubeImportRequest, session: SessionDep):
    try:
        video_id = extract_video_id(request.url)
    except InvalidYouTubeUrlError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    source_url = normalize_source_url(video_id)
    existing_document = session.exec(
        select(Document).where(Document.source_url == source_url)
    ).first()
    if existing_document is not None:
        raise HTTPException(
            status_code=409,
            detail="This YouTube video has already been imported",
        )

    try:
        transcript = download_transcript(request.url)
    except InvalidYouTubeUrlError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except YouTubeVideoUnavailableError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except YouTubeTranscriptUnavailableError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="Could not import YouTube transcript",
        ) from exc

    document = Document(
        filename=transcript.filename,
        file_type="youtube",
        file_path="",
        source_url=transcript.source_url,
        text_content=transcript.text_content,
    )
    session.add(document)
    session.commit()
    session.refresh(document)
    safely_summarize_document(document, session)

    return {
        "status": "saved",
        "document": document,
    }


@app.delete("/documents/{document_id}")
def delete_document(document_id: int, session: SessionDep):
    document = session.get(Document, document_id)

    if document is None:
        raise HTTPException(status_code=404, detail="Document not found")

    file_path = ROOT_DIR / document.file_path

    if file_path.exists():
        file_path.unlink()

    document_id = document.id
    session.delete(document)
    session.commit()
    safely_delete_embedding("document", document_id)

    return {
        "status": "deleted",
        "message": "Document deleted successfully",
    }


@app.put("/notes/{note_id}")
def update_note(note_id: int, updated_note: NoteCreate, session: SessionDep):
    note = session.get(Note, note_id)
    if note is None:
        raise HTTPException(status_code=404, detail="Note not found")

    note.title = updated_note.title
    note.content = updated_note.content
    note.tags = updated_note.tags
    note.updated_at = datetime.now()
    session.add(note)
    session.commit()
    session.refresh(note)
    safely_update_note_embedding(note)

    return {
        "status": "updated",
        "note": note,
    }


@app.delete("/notes/{note_id}")
def delete_note(note_id: int, session: SessionDep):
    note = session.get(Note, note_id)
    if note is None:
        raise HTTPException(status_code=404, detail="Note not found")

    note_id = note.id
    session.delete(note)
    session.commit()
    safely_delete_embedding("note", note_id)

    return {
        "status": "deleted",
        "message": "Note deleted successfully",
    }


@app.post("/collections")
def create_collection(collection: CollectionCreate, session: SessionDep):
    new_collection = Collection(
        name=collection.name,
        description=collection.description,
    )
    session.add(new_collection)
    session.commit()
    session.refresh(new_collection)

    return {
        "status": "saved",
        "collection": new_collection,
    }


@app.get("/collections")
def get_collections(session: SessionDep):
    return session.exec(select(Collection)).all()


@app.get("/collections/{collection_id}")
def get_collection(collection_id: int, session: SessionDep):
    collection = session.get(Collection, collection_id)
    if collection is None:
        raise HTTPException(status_code=404, detail="Collection not found")

    return collection


@app.put("/collections/{collection_id}")
def update_collection(
    collection_id: int, updated_collection: CollectionCreate, session: SessionDep
):
    collection = session.get(Collection, collection_id)
    if collection is None:
        raise HTTPException(status_code=404, detail="Collection not found")

    collection.name = updated_collection.name
    collection.description = updated_collection.description
    collection.updated_at = datetime.now()
    session.add(collection)
    session.commit()
    session.refresh(collection)

    return {
        "status": "updated",
        "collection": collection,
    }


@app.delete("/collections/{collection_id}")
def delete_collection(collection_id: int, session: SessionDep):
    collection = session.get(Collection, collection_id)
    if collection is None:
        raise HTTPException(status_code=404, detail="Collection not found")

    session.delete(collection)
    session.commit()

    return {
        "status": "deleted",
        "message": "Collection deleted successfully",
    }


@app.post("/collections/{collection_id}/notes/{note_id}")
def attach_note_to_collection(collection_id: int, note_id: int, session: SessionDep):
    collection = session.get(Collection, collection_id)
    if collection is None:
        raise HTTPException(status_code=404, detail="Collection not found")

    note = session.get(Note, note_id)
    if note is None:
        raise HTTPException(status_code=404, detail="Note not found")

    existing_link = session.exec(
        select(CollectionNote).where(
            CollectionNote.collection_id == collection_id,
            CollectionNote.note_id == note_id,
        )
    ).first()
    if existing_link is not None:
        return {
            "status": "attached",
            "message": "Note already belongs to collection",
        }

    collection_note = CollectionNote(collection_id=collection_id, note_id=note_id)
    session.add(collection_note)
    session.commit()

    return {
        "status": "attached",
        "message": "Note attached to collection successfully",
    }


@app.delete("/collections/{collection_id}/notes/{note_id}")
def detach_note_from_collection(collection_id: int, note_id: int, session: SessionDep):
    collection = session.get(Collection, collection_id)
    if collection is None:
        raise HTTPException(status_code=404, detail="Collection not found")

    note = session.get(Note, note_id)
    if note is None:
        raise HTTPException(status_code=404, detail="Note not found")

    collection_note = session.exec(
        select(CollectionNote).where(
            CollectionNote.collection_id == collection_id,
            CollectionNote.note_id == note_id,
        )
    ).first()
    if collection_note is None:
        raise HTTPException(status_code=404, detail="Note is not in collection")

    session.delete(collection_note)
    session.commit()

    return {
        "status": "detached",
        "message": "Note detached from collection successfully",
    }


@app.get("/collections/{collection_id}/notes")
def get_collection_notes(collection_id: int, session: SessionDep):
    collection = session.get(Collection, collection_id)
    if collection is None:
        raise HTTPException(status_code=404, detail="Collection not found")

    statement = (
        select(Note)
        .join(CollectionNote, CollectionNote.note_id == Note.id)
        .where(CollectionNote.collection_id == collection_id)
    )

    return session.exec(statement).all()


app.include_router(search_router)
app.include_router(conversations_router)
app.include_router(ai_router)
